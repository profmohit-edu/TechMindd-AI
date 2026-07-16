"""Document ingestion pipeline for RAG knowledge base."""

from __future__ import annotations

import hashlib
import json
import logging
from dataclasses import dataclass
from pathlib import Path

from pypdf import PdfReader

from rag.chunker import TextChunker
from rag.embedder import SentenceTransformerEmbedder
from rag.vector_store import ChunkMetadata, ChromaVectorStore


LOGGER = logging.getLogger("techmindd.rag.ingestion")

_SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".markdown"}


@dataclass(frozen=True)
class IngestionResult:
    """Summary of a completed ingestion run."""

    files: int
    chunks: int


@dataclass(frozen=True)
class SourceDocument:
    """A parsed source document."""

    source: Path
    page: int
    text: str


class IngestionPipeline:
    """Ingest supported documents into Chroma vector store."""

    def __init__(
        self,
        documents_dir: Path = Path("knowledge/documents"),
        embeddings_dir: Path = Path("knowledge/embeddings"),
    ) -> None:
        self._documents_dir = documents_dir
        self._embeddings_dir = embeddings_dir
        self._chunker = TextChunker()
        self._embedder = SentenceTransformerEmbedder()
        self._vector_store = ChromaVectorStore(persist_directory=embeddings_dir)
        self._state_path = embeddings_dir / "ingestion_state.json"

    def ingest(self, documents_path: Path | None = None) -> IngestionResult:
        """Ingest changed documents only; return ingestion summary."""
        target_dir = documents_path or self._documents_dir
        target_dir.mkdir(parents=True, exist_ok=True)
        self._embeddings_dir.mkdir(parents=True, exist_ok=True)

        previous_state = self._load_state()
        next_state: dict[str, str] = {}

        ingested_files = 0
        total_chunks = 0

        all_files = [p for p in sorted(target_dir.rglob("*")) if p.is_file() and p.suffix.lower() in _SUPPORTED_EXTENSIONS]
        print(f"[RAG] Scanning {len(all_files)} supported file(s) in {target_dir}")

        for file_path in all_files:
            source = str(file_path.resolve())
            digest = hashlib.sha256(file_path.read_bytes()).hexdigest()
            next_state[source] = digest

            if previous_state.get(source) == digest:
                LOGGER.debug("Unchanged, skipping: %s", file_path.name)
                continue

            print(f"[RAG] Indexing: {file_path.name}")
            self._vector_store.delete_by_source(source)
            chunks = self._ingest_single_file(file_path)
            total_chunks += chunks
            ingested_files += 1
            LOGGER.info("Ingested %s — %d chunk(s)", file_path.name, chunks)

        removed_sources = set(previous_state).difference(next_state)
        for removed_source in removed_sources:
            self._vector_store.delete_by_source(removed_source)
            LOGGER.info("Removed stale source from index: %s", removed_source)

        self._save_state(next_state)
        return IngestionResult(files=ingested_files, chunks=total_chunks)

    def _ingest_single_file(self, path: Path) -> int:
        """Ingest one file; return number of chunks produced."""
        docs = self._parse_file(path)

        ids: list[str] = []
        texts: list[str] = []
        metadatas: list[ChunkMetadata] = []

        for doc in docs:
            chunks = self._chunker.chunk(doc.text)
            for chunk in chunks:
                ids.append(f"{doc.source.resolve()}::{doc.page}::{chunk.chunk_id}")
                texts.append(chunk.text)
                metadatas.append(
                    ChunkMetadata(
                        filename=doc.source.name,
                        page=doc.page,
                        chunk_id=chunk.chunk_id,
                        source=str(doc.source.resolve()),
                    )
                )

        if not texts:
            return 0

        embeddings = self._embedder.embed_documents(texts)
        self._vector_store.upsert(
            ids=ids,
            documents=texts,
            embeddings=embeddings,
            metadatas=metadatas,
        )
        LOGGER.info("Chunks created for %s: %d", path.name, len(texts))
        return len(texts)

    def _parse_file(self, path: Path) -> list[SourceDocument]:
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            return self._parse_pdf(path)
        if suffix == ".docx":
            return self._parse_docx(path)
        return [
            SourceDocument(
                source=path,
                page=1,
                text=path.read_text(encoding="utf-8", errors="ignore"),
            )
        ]

    def _parse_pdf(self, path: Path) -> list[SourceDocument]:
        reader = PdfReader(str(path))
        docs: list[SourceDocument] = []
        for idx, page in enumerate(reader.pages, start=1):
            docs.append(
                SourceDocument(
                    source=path,
                    page=idx,
                    text=page.extract_text() or "",
                )
            )
        return docs

    def _parse_docx(self, path: Path) -> list[SourceDocument]:
        try:
            import docx  # python-docx
        except ImportError as exc:
            raise ImportError(
                "python-docx is required to parse .docx files. "
                "Install it with: pip install python-docx"
            ) from exc

        doc = docx.Document(str(path))
        text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())
        return [SourceDocument(source=path, page=1, text=text)]

    def _load_state(self) -> dict[str, str]:
        if not self._state_path.exists():
            return {}
        try:
            raw = json.loads(self._state_path.read_text(encoding="utf-8"))
            if not isinstance(raw, dict):
                return {}
            return {str(k): str(v) for k, v in raw.items()}
        except (json.JSONDecodeError, OSError):
            return {}

    def _save_state(self, state: dict[str, str]) -> None:
        self._state_path.write_text(
            json.dumps(state, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
